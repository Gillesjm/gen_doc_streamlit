import streamlit as st
import pandas as pd
import os
import shutil
from docx import Document
from docx.shared import RGBColor
from io import BytesIO

# Fonction de génération du fichier Word
def generate_docx_files(template_path, df):
    """
    Génère un fichier Word personnalisé à partir d'un modèle et d'un DataFrame.
    Retourne un BytesIO contenant le document généré.
    """
    try:
        doc = Document(template_path)
        
        for paragraph in doc.paragraphs:
            for column in df.columns:
                placeholder = f"<<{column}>>"
                if placeholder in paragraph.text:
                    inline_text = paragraph.text.split(placeholder)
                    paragraph.clear()
                    for i, part in enumerate(inline_text):
                        run = paragraph.add_run(part)
                        if i < len(inline_text) - 1:
                            inserted_text = str(df.at[0, column])
                            inserted_text = "non renseigné" if inserted_text == "nan" else inserted_text
                            formatted_run = paragraph.add_run(inserted_text)
                            formatted_run.bold = True
                            formatted_run.font.color.rgb = RGBColor(128, 0, 128)

        # Remplacer dans les tableaux
        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for column in df.columns:
                        placeholder = f"<<{column}>>"
                        if placeholder in cell.text:
                            cell_parts = cell.text.split(placeholder)
                            cell.text = ""
                            for i, part in enumerate(cell_parts):
                                cell_run = cell.paragraphs[0].add_run(part)
                                if i < len(cell_parts) - 1:
                                    inserted_text = str(df.at[0, column])
                                    inserted_text = "non renseigné" if inserted_text == "nan" else inserted_text
                                    formatted_run = cell.paragraphs[0].add_run(inserted_text)
                                    formatted_run.bold = True
                                    formatted_run.font.color.rgb = RGBColor(128, 0, 128)
                                    
        

        # Création du fichier Word en mémoire
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Génération du nom de fichier - Code de la structure + CRV + Date de la dernière visite
        # Vérification et formatage de la date
        try:
            df["Date de la dernière visite"] = pd.to_datetime(df["Date de la dernière visite"], format="%d/%m/%Y", errors="coerce")
            date = df["Date de la dernière visite"].dt.strftime("%Y.%m.%d")[0]  # Format YYYY.MM.DD
        except:
            date = "non_défini"
        
        # Génération du nom de fichier
        output_filename = f"{df.at[0, 'Code de la structure']} CRV {date}.docx"

        return output_filename, output

    except Exception as e:
        st.error(f"Erreur lors de la génération du document : {e}")
        return None

# Interface Streamlit
st.title(":blue[GEN_DOC-📄]")
st.caption("by **Gilles G.**", unsafe_allow_html=False)
# description
# présentation de l'algorithme
st.info("**DESCRIPTION DE L'APPLICATION**"
        "\n - Cette application permet de générer un rapport de visite sous un format Word à partir d'un modèle et d'un fichier Excel."
        "\n - Le fichier Excel est extrait d'Assoconnect et contient les informations de la structure (association) pour remplir le modèle.\n"
        "\n - Le fichier Word généré est personnalisé avec les informations de la structure.\n"
        "\n - le nom du fichier généré est composé du Code de la structure + CRV + Date de la dernière visite.\n"
    )

# Téléchargement du fichier Excel
uploaded_file = st.file_uploader("Téléchargez le fichier Excel d'Assoconnect", type=["xlsx"])

# Traitement du fichier Excel
if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file)
        
        assoc_name = df['Nom'][0]

        # Affichage des colonnes
        st.info("**Aperçu des données du fichier excel**")
        st.success(f"Nom de la structure : **{assoc_name}**")
       
        st.dataframe(df.head())

        # Vérification et conversion des dates
        datetime_cols = df.select_dtypes(include=['datetime64', 'timedelta64']).columns
        for col in datetime_cols:
            df[col] = df[col].dt.strftime('%d/%m/%Y')

        # Sélection du modèle Word
        # chemin complet de l'algorithme
        full_path = os.path.realpath(__file__)
        # nom de l'algorithme
        algo_name = full_path.split('/')[-1]
        path = full_path[:full_path.find(algo_name)]
        #st.write('path =', path)
        template_path = os.path.join(path, "template_assoc.docx")
        
        if not os.path.exists(template_path):
            st.error("Le fichier modèle 'template_assoc.docx' est introuvable. Ajoutez-le au même dossier que l'application.")
        else:
            if st.button("🔄 Générer le document Word"):
                with st.spinner("Génération en cours..."):
                    file_name, word_file = generate_docx_files(template_path, df)
                    if word_file:
                        st.success("✅ Document généré avec succès !")

                        # Permettre le téléchargement du fichier
                        st.download_button(
                            label="📥 Télécharger le fichier Word",
                            data=word_file,
                            file_name=file_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
    except Exception as e:
        st.error(f"Erreur lors du traitement du fichier : {e}")
